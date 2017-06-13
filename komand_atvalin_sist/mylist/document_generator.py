from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

class DocumentCreator():
    def __init__(self):
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        spacing = style.paragraph_format
        spacing.space_after = Pt(0)

        # document.add_heading('Document Title', 0)
        #
        # p = document.add_paragraph('A plain paragraph having some ')

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run('SIA TestDevLab')
        run.add_break()
        run = paragraph.add_run('Valdes loceklim')
        run.add_break()
        run = paragraph.add_run('Ervinam Grinfeldam')

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = paragraph.add_run('Informacijas sistemu testetajs')
        run.add_break()
        run = paragraph.add_run('Edgars Avotins')
        run.add_break()
        run = paragraph.add_run('p.k.111111-11111')
        run.add_break()
        run.add_break()
        run.add_break()
        run = paragraph.add_run(' ')

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run('Iesniegums.')
        run.add_break()
        run.add_break()
        run = paragraph.add_run(' ')

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run("Ludzu pieskirt man ikgadejo apmaksato atvalinajumu no 2015. gada 5. marta lidz 2015. gada 7. martam.")
        paragraph.paragraph_format.first_line_indent = Inches(0.25)


        # p.add_run('bold').bold = True
        # p.add_run(' and some ')
        # p.add_run('italic.').italic = True
        #
        # document.add_heading('Heading, level 1', level=1)
        # document.add_paragraph('Intense quote', style='IntenseQuote')
        #
        # document.add_paragraph(
        #     'first item in unordered list', style='ListBullet'
        # )
        # document.add_paragraph(
        #     'first item in ordered list', style='ListNumber'
        # )
        #
        # document.add_picture('monty-truth.png', width=Inches(1.25))
        #
        # table = document.add_table(rows=1, cols=3)
        # hdr_cells = table.rows[0].cells
        # hdr_cells[0].text = 'Qty'
        # hdr_cells[1].text = 'Id'
        # hdr_cells[2].text = 'Desc'
        # # for item in recordset:
        # #     row_cells = table.add_row().cells
        # #     row_cells[0].text = str(item.qty)
        # #     row_cells[1].text = str(item.id)
        # #     row_cells[2].text = item.desc
        #
        # document.add_page_break()

        document.save('doc_test.docx')